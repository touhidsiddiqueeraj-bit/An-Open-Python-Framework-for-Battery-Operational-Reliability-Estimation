#!/usr/bin/env python3
"""Render manuscript.md to a formatted .docx file using python-docx."""

import os, re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH = os.path.join(SCRIPT_DIR, "manuscript.md")
OUT_PATH = os.path.join(SCRIPT_DIR, "Extension_Paper.docx")
FIG_DIR = os.path.join(SCRIPT_DIR, "figures")

def parse_markdown(md_text):
    """Simple markdown parser for our specific doc structure."""
    lines = md_text.split("\n")
    sections = []
    current_section = None
    current_lines = []
    inside_table = False
    table_lines = []

    def flush():
        nonlocal current_lines, inside_table, table_lines
        if inside_table and table_lines:
            sections.append({"type": "table", "content": table_lines})
            table_lines = []
            inside_table = False
        if current_lines:
            if current_section == "p":
                sections.append({"type": "paragraph", "content": " ".join(current_lines)})
            elif current_section == "li":
                sections.append({"type": "list", "content": current_lines})
            current_lines = []

    for line in lines:
        stripped = line.strip()

        # Detect table
        if stripped.startswith("|") and stripped.endswith("|"):
            if not inside_table:
                flush()
                inside_table = True
                table_lines = [stripped]
            else:
                table_lines.append(stripped)
            continue
        else:
            if inside_table:
                flush()

        # Empty line
        if not stripped:
            flush()
            continue

        # Headings
        if stripped.startswith("---"):
            continue
        h_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if h_match:
            flush()
            level = len(h_match.group(1))
            sections.append({"type": "heading", "level": level, "content": h_match.group(2)})
            continue

        # Ordered list item
        if re.match(r"^\d+\.\s+", stripped):
            current_section = "li"
            current_lines.append(re.sub(r"^\d+\.\s+", "", stripped))
            continue

        # Bullet item
        if stripped.startswith("- "):
            current_section = "li"
            current_lines.append(stripped[2:])
            continue

        # Regular paragraph
        if current_section != "p":
            flush()
        current_section = "p"
        # Remove bold markers for docx (we'll handle bold via runs)
        current_lines.append(stripped)

    flush()
    if inside_table and table_lines:
        sections.append({"type": "table", "content": table_lines})
    return sections


def add_formatted_paragraph(doc, text, bold=False, italic=False, font_size=11,
                            alignment=None, space_after=Pt(6), space_before=Pt(0),
                            font_name="Times New Roman"):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = font_name
    return p


def render_table(sections, doc):
    for sec in sections:
        if sec["type"] == "heading":
            add_formatted_paragraph(
                doc, sec["content"],
                bold=True,
                font_size=14 if sec["level"] == 1 else 12 if sec["level"] == 2 else 11,
                space_before=Pt(12) if sec["level"] == 1 else Pt(8),
                space_after=Pt(6))

        elif sec["type"] == "paragraph":
            t = sec["content"]
            # Handle bold markers **...**
            parts = re.split(r"(\*\*.*?\*\*)", t)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

        elif sec["type"] == "list":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(1)
            for i, item in enumerate(sec["content"]):
                if i > 0:
                    p.add_run("\n").font.size = Pt(11)
                run = p.add_run(f"• {item}")
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

        elif sec["type"] == "table":
            rows_data = []
            for line in sec["content"]:
                cells = [c.strip() for c in line.split("|")[1:-1]]
                if cells:
                    rows_data.append(cells)

            if len(rows_data) >= 2:
                n_cols = max(len(r) for r in rows_data)
                # Align rows
                for i, r in enumerate(rows_data):
                    while len(r) < n_cols:
                        r.append("")

                table = doc.add_table(rows=len(rows_data), cols=n_cols)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for i, row_data in enumerate(rows_data):
                    for j, cell_text in enumerate(row_data):
                        cell = table.cell(i, j)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        run = p.add_run(cell_text)
                        run.font.size = Pt(10)
                        run.font.name = "Times New Roman"
                        if i == 0:  # Header row
                            run.bold = True
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph().paragraph_format.space_after = Pt(4)  # spacing after table


def add_figures(doc):
    figure_files = [
        ("degradation_curves.png", "Battery degradation curves (SOH vs. cycle) for the four NASA classic cells. B0006 crosses the EOL threshold (SOH < 0.70) at approximately cycle 130."),
        ("auc_by_horizon.png", "Model discrimination by prediction horizon. Raw XGBoost AUC is near-random (0.26--0.69) while calibrated AUC is 0.50 across all horizons (isotonic regression preserves rank order)."),
        ("dataset_composition.png", "Dataset composition showing number of discharge cycles per cell (636 total, 64 EOL events)."),
        ("feature_correlation.png", "Feature correlation matrix. SOH and d\\_SOH show moderate correlation with voltage and temperature features."),
        ("ablation.png", "Ablation study showing failure rates across configurations. The 4-cell dataset provides insufficient signal for dispatch differentiation."),
        ("real_vs_synthetic.png", "Comparison of real NASA degradation (left) and synthetic degradation generated by the data pipeline (right). The synthetic generator produces qualitatively similar degradation trajectories."),
        ("scaling_curve.png", "Figure 1: Synthetic scaling study: macro-averaged AUC vs number of batteries (N). Shaded region shows ±1 standard deviation across 3 Monte Carlo seeds. Regime annotations: insufficient (N ≤ 5), marginal (5 < N < 12), reliable (N ≥ 12). Diminishing returns are evident beyond N=12."),
    ]

    for fname, caption in figure_files:
        fpath = os.path.join(FIG_DIR, fname)
        if os.path.exists(fpath):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(fpath, width=Inches(5.5))
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(caption)
            run2.font.size = Pt(9)
            run2.italic = True
            run2.font.name = "Times New Roman"


def main():
    with open(MD_PATH, "r") as f:
        md_text = f.read()

    sections = parse_markdown(md_text)

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    add_formatted_paragraph(
        doc,
        "An Open Python Framework for Battery Operational Reliability Estimation",
        bold=True, font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(4))
    add_formatted_paragraph(
        doc, "Author: Team Dynamic",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(12))

    # Render all sections
    render_table(sections, doc)

    # Insert figures
    doc.add_page_break()
    add_formatted_paragraph(doc, "Figures", bold=True, font_size=14,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after=Pt(12))
    add_figures(doc)

    doc.save(OUT_PATH)
    print(f"Paper rendered to: {OUT_PATH}")
    print(f"Size: {os.path.getsize(OUT_PATH) // 1024} KB")


if __name__ == "__main__":
    main()
